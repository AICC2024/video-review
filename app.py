
import threading
import base64
from dotenv import load_dotenv
import os
from pathlib import Path
load_dotenv(dotenv_path=Path(__file__).resolve().parent / ".env")
import time
from flask import Flask, request, jsonify, send_from_directory, send_file, abort, render_template
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from datetime import datetime
import pytz
import json
from sqlalchemy.ext.mutable import MutableDict
from werkzeug.security import generate_password_hash, check_password_hash
import secrets
from werkzeug.utils import secure_filename, safe_join
from docx import Document
import boto3
from botocore.exceptions import BotoCoreError, ClientError

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})
app.config['MAX_CONTENT_LENGTH'] = 500 * 1024 * 1024  # Allow uploads up to 500MB



# --- Transcript Route ---
@app.route("/transcript/<video_id>", methods=["GET"])
def get_transcript_for_video(video_id):
    try:
        filename = f"transcripts/{video_id}.json"
        if os.path.exists(filename):
            with open(filename, "r") as f:
                data = json.load(f)
            return jsonify(data)
        else:
            return jsonify({"error": "Transcript not found"}), 404
    except Exception as e:
        print("[❌] Error reading transcript:", e)
        return jsonify({"error": "Failed to read transcript"}), 500

# --- On Demand Transcript Fallback Route ---
@app.route("/transcript_on_demand/<video_id>", methods=["GET"])
def transcribe_video_on_demand(video_id):
    import tempfile
    from moviepy.editor import VideoFileClip
    from openai import OpenAI
    try:
        s3_key = f"videos/{video_id}.mp4"
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
            s3_client.download_fileobj(S3_BUCKET, s3_key, temp_video)
            video_path = temp_video.name

        audio_clip = VideoFileClip(video_path).audio
        audio_path = video_path.replace(".mp4", ".mp3")
        audio_clip.write_audiofile(audio_path, codec="mp3")

        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        with open(audio_path, "rb") as f:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=f,
                response_format="verbose_json"
            )

        results = []
        for segment in transcript.segments:
            results.append({
                "start": segment.start,
                "end": segment.end,
                "text": segment.text.strip()
            })

        return jsonify(results)
    except Exception as e:
        print("[❌] Error generating on-demand transcript:", e)
        return jsonify({"error": "Failed to transcribe video"}), 500

# PostgreSQL connection string placeholder (update before deployment)
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get("DATABASE_URL", "postgresql://paulminton@localhost:5432/video_review")
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

EXPORT_FOLDER = os.path.join(os.getcwd(), 'exports')
os.makedirs(EXPORT_FOLDER, exist_ok=True)

S3_BUCKET = 'naveon-video-storage'
S3_REGION = 'us-east-1'  # change if your bucket is in a different region

s3_client = boto3.client('s3')

# --- SILAS Video Review Async Endpoint ---
@app.route('/silas/review_video_async', methods=['POST'])
def silas_review_video_async():
    print("[📥] /silas/review_video_async endpoint triggered")
    import requests
    import subprocess
    import tempfile
    import base64
    import ffmpeg
    from PIL import Image
    from io import BytesIO
    from moviepy.editor import VideoFileClip

    data = request.json
    file_url = data.get("file_url")
    media_type = data.get("media_type")
    video_id = data.get("video_id")

    if not file_url or not media_type or not video_id:
        return jsonify({"error": "Missing required fields"}), 400

    def run_async_review():
        print("🚀 Inside SILAS run_async_review thread")
        with app.app_context():
            print(f"[✅] SILAS background thread started for: {video_id}")
            try:
                # Download video
                video_resp = requests.get(file_url)
                if video_resp.status_code != 200:
                    print("❌ Failed to download video")
                    return

                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
                    temp_video.write(video_resp.content)
                    video_path = temp_video.name

                print(f"[📥] Video downloaded and saved to {video_path}")

                # Transcribe with Whisper
                from openai import OpenAI
                client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

                print("[🔈] Extracting audio from video...")
                audio_clip = VideoFileClip(video_path).audio
                audio_path = video_path.replace(".mp4", ".mp3")
                audio_clip.write_audiofile(audio_path, codec="mp3")

                with open(audio_path, "rb") as f:
                    transcript = client.audio.transcriptions.create(model="whisper-1", file=f)
                full_text = transcript.text

                print(f"[📝] Transcript preview: {full_text[:100]}...")

                # Break into 3s segments
                segments = []
                words = full_text.split()
                chunk_size = int(len(words) / 40) or 1
                for i in range(0, len(words), chunk_size):
                    segments.append(" ".join(words[i:i+chunk_size]))

                # Capture frame every 3 seconds
                duration = VideoFileClip(video_path).duration
                timestamps = [int(t) for t in range(0, int(duration), 3)]

                print(f"[🎞️] Processing {len(timestamps)} frames at 3s intervals")

                for i, ts in enumerate(timestamps):
                    try:
                        frame_path = video_path.replace(".mp4", f"_{ts}.jpg")
                        (
                            ffmpeg
                            .input(video_path, ss=ts)
                            .output(frame_path, vframes=1)
                            .run(quiet=True, overwrite_output=True)
                        )

                        with open(frame_path, "rb") as img_file:
                            img_b64 = base64.b64encode(img_file.read()).decode("utf-8")

                        vision_prompt = [
                            {
                                "type": "text",
                                "text": f"Please review this video scene (Timestamp: {ts}s). Here is the narration: “{segments[i]}” Provide 1–2 visual improvement suggestions for the scene shown."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{img_b64}"
                                }
                            }
                        ]

                        response = client.chat.completions.create(
                            model="gpt-4o",
                            messages=[
                                {
                                    "role": "system",
                                    "content": get_instruction("video")
                                },
                                {
                                    "role": "user",
                                    "content": vision_prompt
                                }
                            ],
                            max_tokens=500
                        )

                        feedback = response.choices[0].message.content.strip()
                        comment = Comment(
                            video_id=video_id,
                            timestamp=str(ts),
                            comment=feedback + "\n\n-- SILAS (Video Review)",
                            user="SILAS"
                        )
                        db.session.add(comment)
                        db.session.commit()
                        print(f"[✅] Saved comment for {ts}s")
                    except Exception as frame_err:
                        print(f"❌ Error processing timestamp {ts}: {frame_err}")
            except Exception as e:
                import traceback
                print("[❌] SILAS video review error:", str(e))
                traceback.print_exc()

    print("✅ Review thread dispatched")
    threading.Thread(target=run_async_review).start()
    return jsonify({"status": "SILAS video review started"}), 202

# --- Admin List S3 Files Route ---
@app.route('/admin/list', methods=['GET'])
def list_s3_files():
    category = request.args.get("category")
    if not category:
        return jsonify({"error": "Missing category"}), 400

    prefix = f"{category}/"
    try:
        response = s3_client.list_objects_v2(Bucket=S3_BUCKET, Prefix=prefix)
        files = [
            obj["Key"].split("/", 1)[-1]
            for obj in response.get("Contents", [])
            if obj["Key"].startswith(prefix) and obj["Key"] != prefix
        ]
        return jsonify(files)
    except Exception as e:
        print(f"[❌] Failed to list {category} files:", str(e))
        return jsonify([]), 500

# --- Archive S3 File Route (Soft Delete) ---
@app.route('/admin/archive', methods=['POST'])
def archive_s3_file():
    data = request.json
    category = data.get("category")
    filename = data.get("filename")

    if not category or not filename:
        return jsonify({'error': 'Missing category or filename'}), 400

    original_key = f"{category}/{filename}"
    archive_key = f"archive/{category}/{filename}"

    try:
        # Copy the file to the archive location
        s3_client.copy_object(
            Bucket=S3_BUCKET,
            CopySource={'Bucket': S3_BUCKET, 'Key': original_key},
            Key=archive_key
        )
        # Delete the original
        s3_client.delete_object(Bucket=S3_BUCKET, Key=original_key)
        return jsonify({'status': 'archived', 'from': original_key, 'to': archive_key})
    except Exception as e:
        print(f"[❌] Failed to archive {original_key}:", str(e))
        return jsonify({'error': 'Failed to archive file'}), 500

db = SQLAlchemy(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    token = db.Column(db.String(64), unique=True, index=True)

class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    video_id = db.Column(db.String(120), nullable=False)
    timestamp = db.Column(db.String(10), nullable=False)
    comment = db.Column(db.Text, nullable=False)
    user = db.Column(db.String(120), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    reactions = db.Column(MutableDict.as_mutable(db.JSON), default=dict)
    page = db.Column(db.Integer, nullable=True)

# --- SlidePage model for storing full text of each storyboard page ---
class SlidePage(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    video_id = db.Column(db.String(120), nullable=False, index=True)
    page_number = db.Column(db.Integer, nullable=False)
    content = db.Column(db.Text, nullable=False)


# User registration route
@app.route('/register', methods=['POST'])
def register():
    data = request.json
    if User.query.filter_by(username=data['username']).first():
        return jsonify({'error': 'Username already exists'}), 400
    hashed_pw = generate_password_hash(data['password'])
    token = secrets.token_hex(32)
    user = User(username=data['username'], password_hash=hashed_pw, token=token)
    db.session.add(user)
    db.session.commit()
    return jsonify({'status': 'registered', 'token': token, 'username': user.username})

# User login route
@app.route('/login', methods=['POST'])
def login():
    data = request.json
    user = User.query.filter_by(username=data['username']).first()
    if not user or not check_password_hash(user.password_hash, data['password']):
        return jsonify({'error': 'Invalid credentials'}), 401
    return jsonify({'token': user.token, 'username': user.username})

@app.route('/comments', methods=['POST'])
def add_comment():
    data = request.json
    token = request.headers.get('Authorization')
    user = User.query.filter_by(token=token).first()
    username = user.username if user else request.json.get("user", "Anonymous")

    comment = Comment(
        video_id=data['video_id'],
        timestamp=data['timestamp'],
        comment=data['comment'],
        user=username,
        page=data.get("page")
    )
    db.session.add(comment)
    db.session.commit()
    return jsonify({'status': 'success'})

@app.route('/comments/<video_id>', methods=['GET', 'OPTIONS'])
def get_comments(video_id):
    comments = Comment.query.filter_by(video_id=video_id).order_by(Comment.timestamp).all()
    return jsonify([
        {
            "id": c.id,
            "timestamp": c.timestamp,
            "comment": c.comment,
            "user": c.user,
            "created_at": c.created_at.isoformat(),
            "reactions": json.loads(c.reactions) if isinstance(c.reactions, str) else (c.reactions or {}),
            "page": c.page,
        }
        for c in comments
    ])

# Route to get unique video_ids from the comments table
@app.route('/comments/unique_video_ids', methods=['GET'])
def get_unique_video_ids():
    try:
        results = db.session.query(Comment.video_id).distinct().all()
        unique_ids = [row[0] for row in results if row[0]]
        return jsonify(unique_ids)
    except Exception as e:
        print("Error fetching unique video_ids:", e)
        return jsonify([]), 500


@app.after_request
def after_request(response):
    origin = request.headers.get("Origin")
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS,PUT,DELETE,PATCH"
    return response

@app.route("/comments", methods=["OPTIONS"])
def comments_options():
    response = jsonify({"status": "ok"})
    origin = request.headers.get("Origin")
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS,PUT,DELETE,PATCH"
    return response


@app.route("/comments/<int:comment_id>", methods=["OPTIONS"])
def comment_options(comment_id):
    response = jsonify({"status": "ok"})
    origin = request.headers.get("Origin")
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS,PUT,DELETE,PATCH"
    return response

@app.route('/comments/<int:comment_id>', methods=['PUT'])
def update_comment(comment_id):
    data = request.json
    comment = Comment.query.get_or_404(comment_id)
    comment.comment = data.get('comment', comment.comment)
    db.session.commit()
    return jsonify({'status': 'updated', 'id': comment.id})


# Route for updating comment reactions
@app.route("/comments/<int:comment_id>/reactions", methods=["OPTIONS"])
def comment_reactions_options(comment_id):
    response = jsonify({"status": "ok"})
    origin = request.headers.get("Origin")
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS,PUT,DELETE,PATCH"
    return response

@app.route('/comments/<int:comment_id>/reactions', methods=['PATCH'])
def update_reactions(comment_id):
    data = request.json
    token = request.headers.get('Authorization')
    user = User.query.filter_by(token=token).first()
    if not user:
        return jsonify({'error': 'Unauthorized'}), 401

    comment = Comment.query.get_or_404(comment_id)

    if comment.reactions is None:
        comment.reactions = {}

    for reaction in data:
        val = comment.reactions.get(reaction, [])
        if isinstance(val, list):
            users = val
        elif isinstance(val, str):
            users = [val]
        else:
            users = []
        if user.username in users:
            users.remove(user.username)
        else:
            users.append(user.username)
        comment.reactions[reaction] = users

    db.session.commit()
    return jsonify({'status': 'reaction toggled', 'id': comment_id, 'reactions': comment.reactions})

# DELETE route for deleting a comment
@app.route('/comments/<int:comment_id>', methods=['DELETE'])
def delete_comment(comment_id):
    comment = Comment.query.get_or_404(comment_id)
    db.session.delete(comment)
    db.session.commit()
    return jsonify({'status': 'deleted', 'id': comment_id})

@app.route('/upload', methods=['POST'])
def upload_video():
    file = request.files.get('video') or request.files.get('file')
    if file is None or file.filename == '':
        return jsonify({'error': 'No video file provided'}), 400

    filename = secure_filename(file.filename)

    try:
        s3_client.upload_fileobj(
            file,
            S3_BUCKET,
            filename,
            ExtraArgs={'ContentType': file.content_type}
        )
        s3_url = f"https://{S3_BUCKET}.s3.amazonaws.com/{filename}"
        print(f"[✅] Uploaded to S3: {s3_url}")
        response_data = {
            'status': 'uploaded',
            'filename': filename,
            'url': s3_url
        }
        print(f"[✅] Returning response: {response_data}")
        return jsonify(response_data)

    except (BotoCoreError, ClientError) as e:
        print(f"[❌] S3 Upload failed: {e}")
        return jsonify({'error': 'Upload to S3 failed'}), 500


# Correct upload route definition (ensuring no duplicates)
@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    return send_file(path, mimetype="video/mp4", conditional=True)

@app.route('/export/<video_id>', methods=['GET'])
def export_comments(video_id):
    comments = Comment.query.filter_by(video_id=video_id).order_by(Comment.page.nullslast(), Comment.timestamp).all()

    base_filename = f"{video_id}_v"
    existing_versions = [f for f in os.listdir(EXPORT_FOLDER) if f.startswith(base_filename) and f.endswith(".docx")]
    version = len(existing_versions) + 1
    export_filename = f"{base_filename}{version}.docx"
    export_path = os.path.join(EXPORT_FOLDER, export_filename)

    local_time = datetime.now().strftime('%Y-%m-%d %I:%M %p')  # local time

    doc = Document()
    doc.add_heading(f"Comments for Video: {video_id}", 0)
    doc.add_paragraph(f"Exported on: {local_time}")
    doc.add_paragraph(f"Version: {version}")
    doc.add_paragraph("")

    for c in comments:
        full = c.comment.strip()
        base = full.split("\n\n--")[0]
        additions = full.split("\n\n--")[1:] if "\n\n--" in full else []

        p = doc.add_paragraph()
        if c.page:
            p.add_run(f"Page {c.page}:").bold = True
        else:
            p.add_run(f"{c.timestamp} seconds").bold = True
        p.add_run(f" — {base} ({c.user})")
        for add in additions:
            lines = add.strip().split("\n")
            meta = lines[0]
            body = "\n".join(lines[1:]).strip()
            doc.add_paragraph(f"{body} ({meta})")
        doc.add_paragraph("")

    doc.save(export_path)
    return send_file(export_path, as_attachment=True)



# Admin asset upload route
@app.route('/admin/upload', methods=['POST'])
def admin_upload_asset():
    file = request.files.get('file')
    category = request.form.get('category')  # 'videos', 'storyboards', 'voiceovers'

    if not file or not category:
        return jsonify({'error': 'File and category are required'}), 400

    filename = secure_filename(file.filename)
    s3_key = f"{category}/{filename}"

    try:
        s3_client.upload_fileobj(
            file,
            S3_BUCKET,
            s3_key,
            ExtraArgs={'ContentType': file.content_type}
        )
        s3_url = f"https://{S3_BUCKET}.s3.amazonaws.com/{s3_key}"
        print(f"[✅] Admin uploaded to S3: {s3_url}")
        return jsonify({'status': 'uploaded', 's3_key': s3_key, 'url': s3_url})
    except (BotoCoreError, ClientError) as e:
        print(f"[❌] Admin S3 upload failed: {e}")
        return jsonify({'error': 'Admin upload to S3 failed'}), 500


# --- SILAS System Instruction Admin API ---
@app.route("/admin/instructions", methods=["GET"])
def get_silas_instruction():
    mode = request.args.get("mode")
    if not mode:
        return jsonify({"error": "Missing mode"}), 400

    try:
        with open("silas_instructions.json", "r") as f:
            data = json.load(f)
        return jsonify({"mode": mode, "content": data.get(mode, "")})
    except Exception as e:
        print("[❌] Failed to load instructions:", e)
        return jsonify({"error": "Unable to load instructions"}), 500
# Route for listing S3 files by category
@app.route('/media', methods=['GET'])
def list_media_by_type():
    category = request.args.get('type')  # 'videos', 'storyboards', 'voiceovers'
    if not category:
        return jsonify({'error': 'Missing type query parameter'}), 400

    try:
        response = s3_client.list_objects_v2(Bucket=S3_BUCKET, Prefix=f"{category}/")
        files = []
        for obj in response.get('Contents', []):
            key = obj['Key']
            if key.endswith('/'):
                continue  # skip folder entries
            filename = key.split('/')[-1]
            file_url = f"https://{S3_BUCKET}.s3.amazonaws.com/{key}"
            files.append({'filename': filename, 'url': file_url})

        return jsonify(files)
    except Exception as e:
        print(f"[❌] Failed to list {category} from S3:", str(e))
        return jsonify({'error': 'Failed to list media files'}), 500


# --- SILAS AI Review Endpoint ---
@app.route('/silas/review', methods=['POST'])
def silas_review():
    """
    Accepts a PDF or DOCX URL, downloads it, extracts the content, sends to SILAS for review,
    and stores returned comments.
    """
    import requests
    import fitz  # PyMuPDF
    data = request.json
    file_url = data.get("file_url")
    media_type = data.get("media_type")
    video_id = data.get("video_id")

    if not file_url or not media_type or not video_id:
        return jsonify({"error": "Missing required fields"}), 400

    # DOCX handling (must come before PDF check)
    if file_url.lower().endswith(".docx"):
        try:
            from openai import OpenAI
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

            s3_key = f"documents/{video_id}.docx"
            temp_path = f"/tmp/{video_id}.docx"
            s3_client.download_file(S3_BUCKET, s3_key, temp_path)

            doc = Document(temp_path)
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            doc_text = "\n\n".join(paragraphs)

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": get_instruction("document")
                    },
                    {
                        "role": "user",
                        "content": f"Here is the content of the document titled {video_id}:\n\n{doc_text}\n\nPlease provide a detailed review with feedback."
                    }
                ],
                max_tokens=1000
            )

            reply = response.choices[0].message.content.strip()
            comment = Comment(
                video_id=video_id,
                timestamp="0",
                comment=reply + "\n\n-- SILAS (Document Review)",
                user="SILAS"
            )
            db.session.add(comment)
            db.session.commit()
            return jsonify({"status": "SILAS docx review completed", "comments_added": 1})
        except Exception as docx_err:
            print("[❌] DOCX SILAS review error:", docx_err)
            return jsonify({"error": "DOCX review failed"}), 500

    # Only handle PDFs for now (storyboards)
    if not file_url.lower().endswith(".pdf"):
        return jsonify({"error": "Only PDF review is supported in this endpoint"}), 400

    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        # Download PDF
        resp = requests.get(file_url)
        if resp.status_code != 200:
            return jsonify({"error": "Failed to download PDF"}), 400

        # Load PDF into PyMuPDF
        pdf_bytes = resp.content
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        num_pages = doc.page_count

        comments_added = 0
        for page_num in range(num_pages):
            page = doc.load_page(page_num)
            page_text = page.get_text().strip()

            # Save the page text for reference but always use Vision for review
            existing_page = SlidePage.query.filter_by(video_id=video_id, page_number=page_num+1).first()
            if existing_page:
                existing_page.content = page_text
            else:
                db.session.add(SlidePage(
                    video_id=video_id,
                    page_number=page_num+1,
                    content=page_text
                ))

            # Run GPT-4o Vision review regardless of text content
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            img_b64 = base64.b64encode(img_bytes).decode("utf-8")

            vision_prompt = [
                {
                    "type": "text",
                    "text": f"Please review this storyboard slide (Page {page_num + 1}). Provide only 2–3 specific, visual improvements. Base your feedback on what you clearly see in the slide and its narration. Avoid vague language like 'if not already present' and do not include Overall Tone or What Works unless explicitly instructed."
                },
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:image/png;base64,{img_b64}"
                    }
                }
            ]

            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {
                        "role": "system",
                        "content": get_instruction("pdf")
                    },
                    {
                        "role": "user",
                        "content": vision_prompt
                    }
                ],
                max_tokens=1000
            )

            raw_reply = response.choices[0].message.content.strip()
            stripped_reply = raw_reply.replace("**", "")
            formatted_reply = f"Slide {page_num + 1}: {stripped_reply}"

            new_comment = Comment(
                video_id=video_id,
                page=page_num + 1,
                timestamp="0",
                comment=formatted_reply + "\n\n-- SILAS (Vision Review)",
                user="SILAS"
            )
            db.session.add(new_comment)
            db.session.commit()
            comments_added += 1
        return jsonify({"status": "SILAS review completed", "comments_added": comments_added, "pages_reviewed": num_pages})
    except Exception as e:
        print("SILAS error:", str(e))
        return jsonify({"error": "SILAS review failed"}), 500


# --- SILAS Chat Endpoint ---
@app.route('/silas/chat', methods=['POST'])
def silas_chat():
    data = request.json
    message = data.get("message", "").strip()
    file_url = data.get("file_url")
    media_type = data.get("media_type")
    video_id = data.get("video_id")

    if not message:
        return jsonify({"error": "Missing message"}), 400

    try:
        from openai import OpenAI
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        chat_image = data.get("chat_image")
        # Load all comments for the video and format them
        all_comments = []
        comment_context = ""
        page_context = ""
        if video_id:
            all_comments = Comment.query.filter_by(video_id=video_id).order_by(Comment.id.asc()).all()
            comment_context = "\n".join([
                f"{c.timestamp or '0:00'} - {c.user}: {c.comment}" for c in all_comments
            ])
            # Load DOCX or storyboard content
            if file_url and file_url.lower().endswith(".docx"):
                try:
                    s3_key = f"documents/{video_id}.docx"
                    temp_path = f"/tmp/{video_id}.docx"
                    s3_client.download_file(S3_BUCKET, s3_key, temp_path)

                    doc = Document(temp_path)
                    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                    page_context = "\n\n".join(paragraphs)
                except Exception as e:
                    print("[⚠️] Failed to extract DOCX for chat:", e)
                    page_context = ""
            else:
                all_pages = SlidePage.query.filter_by(video_id=video_id).order_by(SlidePage.page_number.asc()).all()
                if all_pages:
                    page_context = "\n\n".join([
                        f"Page {p.page_number}:\n{p.content.strip()}" for p in all_pages
                    ])

        if chat_image:
            prompt = (
                "You have been provided an uploaded image. The user may ask about the image's contents. "
                "Please prioritize the image over other document context if questions reference it directly.\n\n"
                f"User Question: {message}\n\n"
                f"Storyboard Pages (if relevant):\n{page_context}\n\nComments:\n{comment_context}"
            )
        else:
            prompt = (
                f"You are SILAS. The user has a question about this {media_type}.\n"
                f"File: {file_url}\n\n"
                f"Storyboard pages:\n{page_context}\n\n"
                f"Comments so far:\n{comment_context}\n\n"
                f"Question: {message}"
            )

        # Check if message references a specific page and prepare image prompt if needed
        import re
        img_prompt = None
        page_match = re.search(r"\bpage (\d{1,2})\b", message, re.IGNORECASE)
        if page_match and file_url and file_url.lower().endswith(".pdf"):
            try:
                import fitz
                import requests
                page_index = int(page_match.group(1)) - 1
                resp = requests.get(file_url)
                if resp.status_code == 200:
                    doc = fitz.open(stream=resp.content, filetype="pdf")
                    if 0 <= page_index < doc.page_count:
                        page = doc.load_page(page_index)
                        pix = page.get_pixmap(dpi=150)
                        img_bytes = pix.tobytes("png")
                        img_b64 = base64.b64encode(img_bytes).decode("utf-8")
                        img_prompt = {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{img_b64}",
                                "detail": "auto"
                            }
                        }
            except Exception as e:
                print(f"[⚠️] Failed to attach page image to chat prompt: {e}")

        # Prioritize PDF page reference over uploaded image
        if page_match and file_url and file_url.lower().endswith(".pdf"):
            chat_image = None  # Discard screenshot if user asked about a page
        elif chat_image and chat_image.startswith("data:image/"):
            img_prompt = {
                "type": "image_url",
                "image_url": {
                    "url": chat_image,
                    "detail": "auto"
                }
            }

        chat_response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": get_instruction("chat")
                },
                {
                    "role": "user",
                    "content": [{"type": "text", "text": prompt}] + ([img_prompt] if img_prompt else [])
                }
            ]
        )

        reply = chat_response.choices[0].message.content.strip()
        return jsonify({"response": reply})

    except Exception as e:
        print("SILAS chat error:", str(e))
        return jsonify({"error": "SILAS chat failed"}), 500
@app.route('/silas/review_async', methods=['POST'])
def silas_review_async():
    import requests
    import fitz
    data = request.json
    file_url = data.get("file_url")
    media_type = data.get("media_type")
    video_id = data.get("video_id")

    if not file_url or not media_type or not video_id:
        return jsonify({"error": "Missing required fields"}), 400

    def run_async_review():
        with app.app_context():
            try:
                # Use OpenAI client (not openai module directly)
                from openai import OpenAI
                client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
                resp = requests.get(file_url)
                if resp.status_code == 200:
                    doc = fitz.open(stream=resp.content, filetype="pdf")
                    for page_num in range(len(doc)):
                        try:
                            page = doc.load_page(page_num)
                            page_text = page.get_text().strip()

                            existing_page = SlidePage.query.filter_by(video_id=video_id, page_number=page_num+1).first()
                            if existing_page:
                                existing_page.content = page_text
                            else:
                                db.session.add(SlidePage(
                                    video_id=video_id,
                                    page_number=page_num+1,
                                    content=page_text
                                ))

                            pix = page.get_pixmap(dpi=150)
                            img_bytes = pix.tobytes("png")
                            img_b64 = base64.b64encode(img_bytes).decode("utf-8")

                            vision_prompt = [
                                {
                                    "type": "text",
                                    "text": f"Please review this storyboard slide (Page {page_num + 1}). Provide only 2–3 specific, visual improvements. Base your feedback on what you clearly see in the slide and its narration. Avoid vague language like 'if not already present' and do not include Overall Tone or What Works unless explicitly instructed."
                                },
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{img_b64}",
                                        "detail": "auto"
                                    }
                                }
                            ]

                            response = client.chat.completions.create(
                                model="gpt-4o",
                                messages=[
                                    {
                                        "role": "system",
                                        "content": (
                                            "You are SILAS, a helpful and supportive assistant that reviews educational media. "
                                            "You may receive an uploaded image and/or references to specific page numbers. "
                                            "If the user mentions 'this image', assume they are referring to an uploaded screenshot. "
                                            "If they reference a page number (e.g., 'page 3'), use that page from the storyboard PDF instead. "
                                            "If both an image and a page are present, prioritize based on what the user clearly refers to. "
                                            "Always clarify your reference in your reply (e.g., 'Based on page 2' or 'In the uploaded image'). "
                                            "If the context is unclear, ask the user a clarifying question before answering."
                                        )
                                    },
                                    {
                                        "role": "user",
                                        "content": vision_prompt
                                    }
                                ],
                                max_tokens=1000
                            )

                            raw_reply = response.choices[0].message.content.strip()
                            stripped_reply = raw_reply.replace("**", "")
                            formatted_reply = f"Slide {page_num + 1}: {stripped_reply}"

                            new_comment = Comment(
                                video_id=video_id,
                                page=page_num + 1,
                                timestamp="0",
                                comment=formatted_reply + "\n\n-- SILAS (Vision Review)",
                                user="SILAS"
                            )
                            db.session.add(new_comment)
                            db.session.commit()
                        except Exception as page_err:
                            print(f"[❌] Error processing page {page_num + 1}: {page_err}")
            except Exception as e:
                print("[❌] SILAS async thread error:", e)

    threading.Thread(target=run_async_review).start()
    return jsonify({"status": "SILAS review started"}), 202


# --- Notify Team Route ---
@app.route("/notify_team", methods=["POST"])
def notify_team():
    data = request.json
    video_id = data.get("video_id")
    reviewer = data.get("reviewer")
    asset_url = data.get("asset_url")
    to_addresses = data.get("to")

    if not video_id or not reviewer or not asset_url or not to_addresses:
        return jsonify({"error": "Missing required fields"}), 400

    try:
        subject = f"Review Complete: {video_id}"
        body = (
            f"{reviewer} has completed their review of {video_id}.\n\n"
            f"Message:\n{data.get('message', '[No message provided]')}\n\n"
            f"You can view the comments and feedback at:\n{asset_url}"
        )

        ses_client = boto3.client("ses", region_name=S3_REGION)
        ses_client.send_email(
            Source="support@naveonguides.com",
            Destination={"ToAddresses": to_addresses},
            Message={
                "Subject": {"Data": subject},
                "Body": {"Text": {"Data": body}},
            }
        )
        return jsonify({"status": "Notification sent"}), 200

    except Exception as e:
        print("[❌] Failed to notify team:", e)
        return jsonify({"error": "Notification failed"}), 500


# --- Notify Comment Route ---
@app.route("/notify_comment", methods=["POST"])
def notify_comment():
    data = request.json
    video_id = data.get("video_id")
    page = data.get("page")
    comment_text = data.get("comment_text")
    reviewer = data.get("reviewer")
    to_addresses = data.get("to")  # list of emails
    asset_url = data.get("asset_url")

    if not video_id or not comment_text or not reviewer or not to_addresses:
        return jsonify({"error": "Missing required fields"}), 400

    try:
        subject = f"@Notify from {reviewer} - Comment on {video_id}"
        page_info = f"Slide {page}" if page else "Timeline Comment"
        body = (
            f"{reviewer} tagged you in a comment on {page_info} of {video_id}:\n\n"
            f"\"{comment_text.strip()}\"\n\n"
            f"View the full review:\n{asset_url}"
        )

        ses_client = boto3.client("ses", region_name=S3_REGION)
        ses_client.send_email(
            Source="support@naveonguides.com",
            Destination={"ToAddresses": to_addresses},
            Message={
                "Subject": {"Data": subject},
                "Body": {"Text": {"Data": body}},
            }
        )
        return jsonify({"status": "Comment notification sent"}), 200

    except Exception as e:
        print("[❌] Failed to send comment notification:", e)
        return jsonify({"error": "Failed to notify"}), 500


# Serve the admin instruction editor UI
@app.route("/admin/instructions-editor")
def serve_instruction_editor():
    return render_template("admin_instructions.html")


@app.route("/admin/instructions", methods=["POST"])
def save_silas_instruction():
    data = request.json
    mode = data.get("mode")
    content = data.get("content")
    if not mode or content is None:
        return jsonify({"error": "Missing mode or content"}), 400

    try:
        path = "silas_instructions.json"
        instructions = {}
        if os.path.exists(path):
            with open(path, "r") as f:
                instructions = json.load(f)

        instructions[mode] = content
        with open(path, "w") as f:
            json.dump(instructions, f, indent=2)

        return jsonify({"status": "saved", "mode": mode})
    except Exception as e:
        print("[❌] Failed to save instructions:", e)
        return jsonify({"error": "Unable to save instructions"}), 500
@app.route("/docx_text/<video_id>", methods=["GET"])
def extract_docx_text(video_id):
    try:
        s3_key = f"documents/{video_id}.docx"
        temp_path = f"/tmp/{video_id}.docx"
        s3_client.download_file(S3_BUCKET, s3_key, temp_path)

        doc = Document(temp_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return jsonify({ "content": "\n\n".join(paragraphs) })
    except Exception as e:
        print("[❌] Failed to extract DOCX text:", e)
        return jsonify({ "error": str(e) }), 500