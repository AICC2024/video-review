from flask import Flask, request, jsonify, send_from_directory, send_file, abort
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from datetime import datetime
import pytz
import json
from sqlalchemy.ext.mutable import MutableDict
from werkzeug.security import generate_password_hash, check_password_hash
import secrets
import os
from werkzeug.utils import secure_filename, safe_join
from docx import Document

app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "*"}})

# PostgreSQL connection string placeholder (update before deployment)
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get("DATABASE_URL", "postgresql://paulminton@localhost:5432/video_review")
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

UPLOAD_FOLDER = os.path.join(os.getcwd(), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

EXPORT_FOLDER = os.path.join(os.getcwd(), 'exports')
os.makedirs(EXPORT_FOLDER, exist_ok=True)

db = SQLAlchemy(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(256), nullable=False)
    token = db.Column(db.String(64), unique=True, index=True)

class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    video_id = db.Column(db.String(120), nullable=False)
    timestamp = db.Column(db.String(10), nullable=False)
    comment = db.Column(db.Text, nullable=False)
    user = db.Column(db.String(120), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    reactions = db.Column(MutableDict.as_mutable(db.JSON), default=dict)  # new line added


# User registration route
@app.route('/register', methods=['POST'])
def register():
    data = request.json
    if User.query.filter_by(username=data['username']).first():
        return jsonify({'error': 'Username already exists'}), 400
    hashed_pw = generate_password_hash(data['password'])
    token = secrets.token_hex(32)
    user = User(username=data['username'], password_hash=hashed_pw, token=token)
    db.session.add(user)
    db.session.commit()
    return jsonify({'status': 'registered', 'token': token, 'username': user.username})

# User login route
@app.route('/login', methods=['POST'])
def login():
    data = request.json
    user = User.query.filter_by(username=data['username']).first()
    if not user or not check_password_hash(user.password_hash, data['password']):
        return jsonify({'error': 'Invalid credentials'}), 401
    return jsonify({'token': user.token, 'username': user.username})

@app.route('/comments', methods=['POST'])
def add_comment():
    data = request.json
    token = request.headers.get('Authorization')
    user = User.query.filter_by(token=token).first()
    username = user.username if user else request.json.get("user", "Anonymous")

    comment = Comment(
        video_id=data['video_id'],
        timestamp=data['timestamp'],
        comment=data['comment'],
        user=username
    )
    db.session.add(comment)
    db.session.commit()
    return jsonify({'status': 'success'})

@app.route('/comments/<video_id>', methods=['GET', 'OPTIONS'])
def get_comments(video_id):
    comments = Comment.query.filter_by(video_id=video_id).order_by(Comment.timestamp).all()
    return jsonify([
        {
            "id": c.id,
            "timestamp": c.timestamp,
            "comment": c.comment,
            "user": c.user,
            "created_at": c.created_at.isoformat(),
            "reactions": json.loads(c.reactions) if isinstance(c.reactions, str) else (c.reactions or {})
        }
        for c in comments
    ])


@app.after_request
def after_request(response):
    origin = request.headers.get("Origin")
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    return response

@app.route("/comments", methods=["OPTIONS"])
def comments_options():
    response = jsonify({"status": "ok"})
    origin = request.headers.get("Origin")
    if origin:
        response.headers["Access-Control-Allow-Origin"] = origin
    response.headers["Access-Control-Allow-Headers"] = "Content-Type,Authorization"
    response.headers["Access-Control-Allow-Methods"] = "GET,POST,OPTIONS"
    return response


@app.route('/comments/<int:comment_id>', methods=['PUT'])
def update_comment(comment_id):
    data = request.json
    comment = Comment.query.get_or_404(comment_id)
    comment.comment = data.get('comment', comment.comment)
    db.session.commit()
    return jsonify({'status': 'updated', 'id': comment.id})


# Route for updating comment reactions
@app.route('/comments/<int:comment_id>/reactions', methods=['PATCH'])
def update_reactions(comment_id):
    data = request.json
    token = request.headers.get('Authorization')
    user = User.query.filter_by(token=token).first()
    if not user:
        return jsonify({'error': 'Unauthorized'}), 401

    comment = Comment.query.get_or_404(comment_id)

    if comment.reactions is None:
        comment.reactions = {}

    for reaction in data:
        val = comment.reactions.get(reaction, [])
        if isinstance(val, list):
            users = val
        elif isinstance(val, str):
            users = [val]
        else:
            users = []
        if user.username in users:
            users.remove(user.username)
        else:
            users.append(user.username)
        comment.reactions[reaction] = users

    db.session.commit()
    return jsonify({'status': 'reaction toggled', 'id': comment_id, 'reactions': comment.reactions})

# DELETE route for deleting a comment
@app.route('/comments/<int:comment_id>', methods=['DELETE'])
def delete_comment(comment_id):
    comment = Comment.query.get_or_404(comment_id)
    db.session.delete(comment)
    db.session.commit()
    return jsonify({'status': 'deleted', 'id': comment_id})

@app.route('/upload', methods=['POST'])
def upload_video():
    file = request.files.get('video') or request.files.get('file')
    if file is None:
        return jsonify({'error': 'No video file provided'}), 400
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    print(f"[DEBUG] Saved video to: {filepath}")
    return jsonify({'status': 'uploaded', 'filename': filename})


# Correct upload route definition (ensuring no duplicates)
@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

@app.route('/export/<video_id>', methods=['GET'])
def export_comments(video_id):
    comments = Comment.query.filter_by(video_id=video_id).order_by(Comment.timestamp).all()

    base_filename = f"{video_id}_v"
    existing_versions = [f for f in os.listdir(EXPORT_FOLDER) if f.startswith(base_filename) and f.endswith(".docx")]
    version = len(existing_versions) + 1
    export_filename = f"{base_filename}{version}.docx"
    export_path = os.path.join(EXPORT_FOLDER, export_filename)

    local_time = datetime.now().strftime('%Y-%m-%d %I:%M %p')  # local time

    doc = Document()
    doc.add_heading(f"Comments for Video: {video_id}", 0)
    doc.add_paragraph(f"Exported on: {local_time}")
    doc.add_paragraph(f"Version: {version}")
    doc.add_paragraph("")

    for c in comments:
        full = c.comment.strip()
        base = full.split("\n\n--")[0]
        additions = full.split("\n\n--")[1:] if "\n\n--" in full else []

        p = doc.add_paragraph()
        p.add_run(f"{c.timestamp} seconds").bold = True
        p.add_run(f" — {base} ({c.user})")
        for add in additions:
            lines = add.strip().split("\n")
            meta = lines[0]
            body = "\n".join(lines[1:]).strip()
            doc.add_paragraph(f"{body} ({meta})")
        doc.add_paragraph("")

    doc.save(export_path)
    return send_file(export_path, as_attachment=True)

if __name__ == '__main__':
    app.run(host="127.0.0.1", debug=True)
